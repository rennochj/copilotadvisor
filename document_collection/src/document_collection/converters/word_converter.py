"""Word to Markdown converter implementation."""

import zipfile
from pathlib import Path
from typing import Any

from document_collection.core.interfaces import DocumentConverter


class WordConverter(DocumentConverter):
    """Convert Word documents to Markdown format with image extraction."""

    def can_convert(self, file_path: Path) -> bool:
        """Check if this converter can handle Word documents."""
        return str(file_path).lower().endswith(".docx")

    def get_supported_formats(self) -> list[str]:
        """Return list of supported formats."""
        return ["docx"]

    async def convert(self, input_path: Path, output_path: Path, **kwargs: Any) -> Path:
        """Convert Word document to Markdown using python-docx with image extraction."""
        try:
            from docx import Document

            # Read Word document
            doc = Document(str(input_path))

            # Create images directory
            images_dir = output_path.parent / "images"
            images_dir.mkdir(parents=True, exist_ok=True)

            # Extract embedded images from the document's media folder
            image_counter = 1
            image_mapping = {}  # Map original image names to new filenames

            try:
                # Access the document as a zip file to extract images
                with zipfile.ZipFile(str(input_path), 'r') as docx_zip:
                    # Look for images in the media directory
                    for file_info in docx_zip.filelist:
                        if file_info.filename.startswith('word/media/'):
                            # Extract image
                            image_data = docx_zip.read(file_info.filename)

                            # Determine file extension
                            original_name = Path(file_info.filename).name
                            file_ext = Path(file_info.filename).suffix.lower()
                            if not file_ext:
                                file_ext = '.png'  # Default

                            # Create new filename
                            new_filename = f"{input_path.stem}_image_{image_counter:03d}{file_ext}"
                            image_path = images_dir / new_filename

                            # Save image
                            with open(image_path, 'wb') as img_file:
                                img_file.write(image_data)

                            # Map original to new filename
                            image_mapping[original_name] = new_filename
                            image_counter += 1

            except Exception:
                # Continue if image extraction fails
                pass

            # Convert document content to markdown
            content_parts = [f"# {input_path.stem}\n"]
            content_parts.append("Converted from Word document\n")

            if image_counter > 1:
                content_parts.append(f"*This document contains {image_counter - 1} extracted images stored in the `images/` directory.*\n")

            # Process paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # Basic formatting detection with null check
                    style_name = getattr(paragraph.style, 'name', None) if paragraph.style else None
                    if style_name and style_name.startswith('Heading'):
                        level = int(style_name.split()[-1]) if style_name.split()[-1].isdigit() else 1
                        content_parts.append(f"{'#' * min(level + 1, 6)} {text}\n")
                    else:
                        content_parts.append(f"{text}\n")

            # Process tables
            for table in doc.tables:
                content_parts.append("\n")
                # Create markdown table
                for i, row in enumerate(table.rows):
                    row_cells = [cell.text.strip() for cell in row.cells]
                    content_parts.append("| " + " | ".join(row_cells) + " |\n")

                    # Add header separator for first row
                    if i == 0:
                        content_parts.append("| " + " | ".join(["---"] * len(row_cells)) + " |\n")
                content_parts.append("\n")

            # Add image references for any images found
            if image_mapping:
                content_parts.append("\n## Extracted Images\n\n")
                for i, (_original, new_filename) in enumerate(image_mapping.items(), 1):
                    content_parts.append(f"![Image {i}](images/{new_filename})\n\n")

            # Write markdown content
            markdown_content = "\n".join(content_parts)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)

            return output_path

        except ImportError:
            # Fallback if python-docx is not available
            output_path.parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(f"# {input_path.stem}\n\nWord conversion requires python-docx library.\n")
            return output_path
        except Exception as e:
            # Create error file
            output_path.parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(f"# {input_path.stem}\n\nError converting Word document: {str(e)}\n")
            return output_path
